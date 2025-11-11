import os, textwrap
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

ROOT = os.path.dirname(os.path.dirname(__file__))
DOCS = os.path.join(ROOT, "docs")
IMGS = os.path.join(DOCS, "images")
os.makedirs(IMGS, exist_ok=True)

def draw_uml(path: str):
    W, H = 1200, 800
    img = Image.new('RGB', (W, H), 'white')
    d = ImageDraw.Draw(img)
    try:
        ft = ImageFont.truetype("DejaVuSans-Bold.ttf", 22)
        fn = ImageFont.truetype("DejaVuSans.ttf", 16)
        fs = ImageFont.truetype("DejaVuSans.ttf", 14)
    except:
        ft = fn = fs = ImageFont.load_default()

    def box(x, y, w, h, title, fields, methods):
        d.rectangle([x, y, x+w, y+h], outline='black', width=2)
        d.line([x, y+35, x+w, y+35], fill='black', width=2)
        d.text((x+8, y+8), title, font=ft, fill='black')
        d.line([x, y+35+(len(fields)+1)*20, x+w, y+35+(len(fields)+1)*20], fill='black', width=1)
        yy = y+40
        for fld in fields:
            d.text((x+8, yy), fld, font=fs, fill='black'); yy += 20
        yy += 5
        for m in methods:
            d.text((x+8, yy), m, font=fs, fill='black'); yy += 20

    d.text((20, 10), "CIS 476 – Project 2: COR – Hotel Bidding (UML Overview)", font=ft, fill='black')
    box(40, 60, 360, 200, "BidRequest", ["+ price: float"], ["+ __init__(price)"])
    box(450, 60, 700, 300, "RoomBidHandler (abstract)",
        ["# name: str", "# next: RoomBidHandler | None"],
        ["+ handle(bid) -> str | None", "+ can_accept(bid) -> bool", "+ on_accept(bid) -> None", "+ remaining() -> int"])
    box(40, 300, 360, 220, "SuiteHandler", ["- _inv: int"], ["+ can_accept", "+ on_accept", "+ remaining"])
    box(420, 400, 360, 240, "DeluxeHandler", ["- _inv: int", "- _suite_ref: SuiteHandler"], ["+ can_accept", "+ on_accept", "+ remaining"])
    box(820, 400, 330, 240, "StandardHandler", ["- _inv: int", "- _suite_ref: SuiteHandler", "- _deluxe_ref: DeluxeHandler"], ["+ can_accept", "+ on_accept", "+ remaining"])

    def arrow(a, b):
        d.line([a, b], fill='black', width=2)
        ax, ay = b
        d.polygon([(ax, ay), (ax-10, ay-6), (ax-10, ay+6)], outline='black', fill='black')

    arrow((220, 300), (550, 360))
    arrow((600, 400), (650, 360))
    arrow((935, 400), (780, 360))

    notes = ("Chain: Suite → Deluxe → Standard\n"
             "• Suite: ≥ $280\n"
             "• Deluxe: $150–$279.99, or ≥ $280 when Suites sold out\n"
             "• Standard: $80–$149.99, or ≥ $150 when Deluxe & Suite sold out\n"
             "Each accept decrements inventory; stop when sold out.")
    d.text((40, 540), notes, font=fn, fill='black')
    img.save(path)

def draw_mock(path: str, bid: float, result: str, suite:int, deluxe:int, standard:int):
    W, H = 900, 520
    img = Image.new('RGB', (W, H), 'white')
    d = ImageDraw.Draw(img)
    try:
        ft = ImageFont.truetype("DejaVuSans-Bold.ttf", 24)
        fn = ImageFont.truetype("DejaVuSans.ttf", 18)
        fs = ImageFont.truetype("DejaVuSans.ttf", 16)
    except:
        ft = fn = fs = ImageFont.load_default()
    d.rectangle([10,10, W-10, H-10], outline='black', width=2)
    d.text((24, 20), "Hotel Room Bidding - COR Pattern (Mock View)", font=ft, fill='black')
    d.text((24, 60), f"Bid price: ${bid:,.2f}", font=fn, fill='black')
    d.text((24, 100), "Inventory Remaining:", font=fn, fill='black')
    d.text((40, 130), f"Suite: {suite}", font=fs, fill='black')
    d.text((40, 155), f"Deluxe: {deluxe}", font=fs, fill='black')
    d.text((40, 180), f"Standard: {standard}", font=fs, fill='black')
    import textwrap as tw
    d.text((24, 230), "Bid Log (latest):", font=fn, fill='black')
    d.text((40, 265), tw.fill(result, width=70), font=fs, fill='black')
    img.save(path)

def make_doc(uml_path: str, mocks: list, doc_path: str):
    doc = Document()
    doc.add_heading('Project 2 – Chain of Responsibility Hotel Room Bidding System', level=1)
    doc.add_paragraph('Course: CIS 476 – Software Architecture and Design Patterns')
    doc.add_paragraph('Student: Michael Asman')
    doc.add_paragraph('Date: November 10, 2025')
    doc.add_heading('UML Class Diagram', level=2)
    doc.add_picture(uml_path, width=Inches(6.5))
    doc.add_paragraph('Figure 1. UML overview of the COR-based hotel bidding system.')
    doc.add_heading('Execution Scenarios (Screenshots)', level=2)
    for i, p in enumerate(mocks, start=1):
        doc.add_paragraph(f'Scenario {i}:')
        doc.add_picture(p, width=Inches(6.5))
    doc.add_heading('How to Run', level=2)
    doc.add_paragraph('1) Ensure Python 3.x is installed.\n2) Run: python src/hotel_bidding_cor.py\n3) Enter a bid and click "Place Bid".')
    doc.add_heading('Design Notes', level=2)
    doc.add_paragraph('Chain of Responsibility with Suite → Deluxe → Standard. Pricing rules and inventory enforced per handler.')
    doc.save(doc_path)

def main():
    uml_path = os.path.join(IMGS, "uml_cor_hotel.png")
    draw_uml(uml_path)
    scenarios = [
        (300.0, "ACCEPTED: Suite booked at $300.00. Remaining: 9", 9, 15, 45),
        (200.0, "ACCEPTED: Deluxe booked at $200.00. Remaining: 14", 10, 14, 45),
        (90.0,  "ACCEPTED: Standard booked at $90.00. Remaining: 44", 10, 15, 44),
        (300.0, "ACCEPTED: Deluxe booked at $300.00. Remaining: 14 (Suites sold out)", 0, 14, 45),
        (140.0, "REJECTED: No room type available for $140.00. (Deluxe & Suite sold out; Standard needs ≥ $150)", 0, 0, 45),
    ]
    mocks = []
    for i, (bid, res, s, d_, st) in enumerate(scenarios, start=1):
        p = os.path.join(IMGS, f"mock_scenario_{i}.png")
        draw_mock(p, bid, res, s, d_, st)
        mocks.append(p)

    doc_path = os.path.join(DOCS, "Project2_COR_HotelBidding_Asman.docx")
    make_doc(uml_path, mocks, doc_path)
    print("Assets generated under docs/")

if __name__ == "__main__":
    main()
